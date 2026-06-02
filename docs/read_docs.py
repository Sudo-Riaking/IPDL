import os
from pypdf import PdfReader
import docx

def extract_pdf(pdf_path):
    print(f"Extracting PDF: {pdf_path}")
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for i, page in enumerate(reader.pages):
            text += f"\n--- Page {i+1} ---\n"
            text += page.extract_text()
        return text
    except Exception as e:
        return f"Error extracting {pdf_path}: {e}"

def extract_docx(docx_path):
    print(f"Extracting DOCX: {docx_path}")
    try:
        doc = docx.Document(docx_path)
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                text += " | ".join(row_text) + "\n"
        return text
    except Exception as e:
        return f"Error extracting {docx_path}: {e}"

def main():
    files = [
        "Rapport_Phase_Inception_UMMISCO_Groupe8.pdf",
        "Compte_rendu_IPDL.pdf",
        "Compte rendu — Notes de séance IPDL.pdf",
        "Phase1_IPDL_UMMISCO (1).docx",
        "Phase 1 du Projet IPDL — La phase d'Inception (Création).docx",
        "Notes.docx",
        "Notes_Hann.docx",
        "notes_questions.docx"
    ]
    
    out_dir = "extracted_text"
    os.makedirs(out_dir, exist_ok=True)
    
    for f in files:
        if not os.path.exists(f):
            print(f"File not found: {f}")
            continue
        
        ext = os.path.splitext(f)[1].lower()
        if ext == ".pdf":
            text = extract_pdf(f)
        elif ext == ".docx":
            text = extract_docx(f)
        else:
            print(f"Unsupported extension: {f}")
            continue
            
        out_name = os.path.join(out_dir, f + ".txt")
        with open(out_name, "w", encoding="utf-8") as out_f:
            out_f.write(text)
        print(f"Saved to {out_name}")

if __name__ == "__main__":
    main()
